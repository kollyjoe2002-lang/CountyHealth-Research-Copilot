from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document

from ai.classifier import classify_question
from ai.executor import execute_plan
from ai.exporter import (
    default_report_filename,
    export_docx_bytes,
    export_markdown_bytes,
)
from ai.figures import (
    build_evidence_figure,
    export_evidence_figure_png,
)
from ai.planner import build_analysis_plan
from ai.report_writer import write_research_report
from ai.resolver import resolve_plan
from ai.validation import validate_question


QUESTIONS = [
    (
        "trend",
        (
            "Show the trend in ischemic heart disease "
            "in Albany County, Wyoming from 2000 to 2019."
        ),
    ),
    (
        "ranking",
        (
            "Which counties had the highest diabetes "
            "YLL rates in 2019?"
        ),
    ),
    (
        "disparity",
        (
            "Compare diabetes burden among Black and "
            "White groups in 2019."
        ),
    ),
    (
        "profile",
        "Tell me about Albany County, Wyoming.",
    ),
]


PNG_SIGNATURE = b"\x89PNG\r\n\x1a\n"


def build_report(
    question: str,
):
    classified = classify_question(
        question
    )

    validate_question(
        classified
    )

    plan = build_analysis_plan(
        classified
    )

    resolved_plan = resolve_plan(
        classified,
        plan,
    )

    if resolved_plan.unresolved_items:
        raise AssertionError(
            "Question unexpectedly had unresolved items: "
            f"{resolved_plan.unresolved_items}"
        )

    evidence = execute_plan(
        resolved_plan
    )

    report = write_research_report(
        evidence
    )

    return evidence, report


def assert_png(
    png_bytes: bytes,
    case_name: str,
) -> None:
    if not png_bytes:
        raise AssertionError(
            f"{case_name}: PNG export was empty."
        )

    if not png_bytes.startswith(
        PNG_SIGNATURE
    ):
        raise AssertionError(
            f"{case_name}: output was not a valid PNG."
        )

    if len(png_bytes) < 1000:
        raise AssertionError(
            f"{case_name}: PNG was suspiciously small: "
            f"{len(png_bytes)} bytes."
        )


def assert_markdown(
    markdown_bytes: bytes,
    report,
    case_name: str,
) -> None:
    if not markdown_bytes:
        raise AssertionError(
            f"{case_name}: Markdown export was empty."
        )

    text = markdown_bytes.decode(
        "utf-8"
    )

    required_fragments = [
        report.title,
        report.question,
        "## Executive Summary",
        "## Methods",
        "## Key Findings",
        "## Limitations",
        "## Analytical Provenance",
    ]

    for fragment in required_fragments:
        if fragment not in text:
            raise AssertionError(
                f"{case_name}: Markdown export "
                f"was missing {fragment!r}."
            )

    for finding in report.findings:
        if finding not in text:
            raise AssertionError(
                f"{case_name}: Markdown export "
                "was missing a validated finding: "
                f"{finding!r}"
            )

    for limitation in report.limitations:
        if limitation not in text:
            raise AssertionError(
                f"{case_name}: Markdown export "
                "was missing a limitation: "
                f"{limitation!r}"
            )


def extract_docx_text(
    docx_bytes: bytes,
) -> str:
    document = Document(
        BytesIO(
            docx_bytes
        )
    )

    paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
    ]

    table_text = []

    for table in document.tables:
        for row in table.rows:
            table_text.extend(
                cell.text
                for cell in row.cells
            )

    return "\n".join(
        [
            *paragraphs,
            *table_text,
        ]
    )


def assert_docx(
    docx_bytes: bytes,
    report,
    case_name: str,
) -> None:
    if not docx_bytes:
        raise AssertionError(
            f"{case_name}: DOCX export was empty."
        )

    if len(docx_bytes) < 5000:
        raise AssertionError(
            f"{case_name}: DOCX was suspiciously small: "
            f"{len(docx_bytes)} bytes."
        )

    with ZipFile(
        BytesIO(
            docx_bytes
        )
    ) as archive:
        names = set(
            archive.namelist()
        )

        if "word/document.xml" not in names:
            raise AssertionError(
                f"{case_name}: DOCX package did not "
                "contain word/document.xml."
            )

    text = extract_docx_text(
        docx_bytes
    )

    required_fragments = [
        report.title,
        report.question,
        "Methods",
        "Key Findings",
        "Limitations",
        "Analytical Provenance",
    ]

    for fragment in required_fragments:
        if fragment not in text:
            raise AssertionError(
                f"{case_name}: DOCX export "
                f"was missing {fragment!r}."
            )

    for finding in report.findings:
        if finding not in text:
            raise AssertionError(
                f"{case_name}: DOCX export "
                "was missing a validated finding: "
                f"{finding!r}"
            )


def assert_filename(
    report,
    extension: str,
    filename: str,
    case_name: str,
) -> None:
    if not filename.endswith(
        f".{extension}"
    ):
        raise AssertionError(
            f"{case_name}: filename had "
            f"unexpected extension: {filename}"
        )

    if " " in filename:
        raise AssertionError(
            f"{case_name}: filename contains spaces: "
            f"{filename}"
        )

    if not filename:
        raise AssertionError(
            f"{case_name}: filename was empty."
        )


def run_case(
    case_name: str,
    question: str,
) -> None:
    print("\n" + "=" * 88)
    print(
        f"CASE: {case_name.upper()}"
    )
    print("=" * 88)

    print(
        "Question:",
        question,
    )

    evidence, report = build_report(
        question
    )

    figure = build_evidence_figure(
        evidence
    )

    if figure is None:
        raise AssertionError(
            f"{case_name}: figure builder returned None."
        )

    if not figure.axes:
        raise AssertionError(
            f"{case_name}: figure had no axes."
        )

    print(
        "Figure object: PASS"
    )

    png_bytes = export_evidence_figure_png(
        evidence
    )

    assert_png(
        png_bytes,
        case_name,
    )

    print(
        f"PNG export: PASS "
        f"({len(png_bytes):,} bytes)"
    )

    markdown_bytes = export_markdown_bytes(
        report
    )

    assert_markdown(
        markdown_bytes,
        report,
        case_name,
    )

    print(
        f"Markdown export: PASS "
        f"({len(markdown_bytes):,} bytes)"
    )

    docx_bytes = export_docx_bytes(
        report,
        include_evidence_tables=True,
        evidence_row_limit=20,
    )

    assert_docx(
        docx_bytes,
        report,
        case_name,
    )

    print(
        f"DOCX export: PASS "
        f"({len(docx_bytes):,} bytes)"
    )

    markdown_filename = (
        default_report_filename(
            report,
            "md",
        )
    )

    docx_filename = (
        default_report_filename(
            report,
            "docx",
        )
    )

    assert_filename(
        report,
        "md",
        markdown_filename,
        case_name,
    )

    assert_filename(
        report,
        "docx",
        docx_filename,
        case_name,
    )

    print(
        "Filename generation: PASS"
    )

    print(
        "Markdown filename:",
        markdown_filename,
    )

    print(
        "DOCX filename:",
        docx_filename,
    )

    print(
        f"{case_name.upper()}: PASS"
    )


def main() -> None:
    print("=" * 88)
    print(
        "EpiCounty V1 Tier 3.2 "
        "Export and Figure Integration Validation"
    )
    print("=" * 88)

    for case_name, question in QUESTIONS:
        run_case(
            case_name,
            question,
        )

    print("\n" + "=" * 88)
    print(
        "TIER 3.2 EXPORT/FIGURE INTEGRATION: "
        "ALL CASES PASSED"
    )
    print("=" * 88)


if __name__ == "__main__":
    main()