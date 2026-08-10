from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[2]

if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


from ai.classifier import classify_question  # noqa: E402
from ai.executor import execute_plan  # noqa: E402
from ai.exporter import (  # noqa: E402
    save_docx_report,
    save_markdown_report,
)
from ai.planner import build_analysis_plan  # noqa: E402
from ai.report_writer import write_research_report  # noqa: E402
from ai.resolver import resolve_plan  # noqa: E402


OUTPUT_DIR = (
    PROJECT_ROOT
    / "outputs"
    / "validation"
)


QUESTION = (
    "Tell me about Albany County, Wyoming."
)


def main() -> None:
    print("=" * 80)
    print("AI Report Export Validation")
    print("=" * 80)

    classified = classify_question(
        QUESTION
    )

    plan = build_analysis_plan(
        classified
    )

    resolved = resolve_plan(
        classified,
        plan,
    )

    evidence = execute_plan(
        resolved
    )

    report = write_research_report(
        evidence
    )

    markdown_path = save_markdown_report(
        report,
        OUTPUT_DIR
        / "albany_county_report.md",
    )

    docx_path = save_docx_report(
        report,
        OUTPUT_DIR
        / "albany_county_report.docx",
        include_evidence_tables=True,
        evidence_row_limit=10,
    )

    if not markdown_path.exists():
        raise AssertionError(
            "Markdown report was not created."
        )

    if not docx_path.exists():
        raise AssertionError(
            "DOCX report was not created."
        )

    if markdown_path.stat().st_size == 0:
        raise AssertionError(
            "Markdown report is empty."
        )

    if docx_path.stat().st_size == 0:
        raise AssertionError(
            "DOCX report is empty."
        )

    document = Document(
        docx_path
    )

    paragraph_text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    required_text = [
        "County Public Health Profile",
        "Research Question",
        "Executive Summary",
        "Methods",
        "Key Findings",
        "Limitations",
        "Analytical Provenance",
        "Mean BMI",
    ]

    for value in required_text:
        if value not in paragraph_text:
            raise AssertionError(
                f"DOCX report is missing: {value}"
            )

    if len(document.tables) < 1:
        raise AssertionError(
            "DOCX report contains no evidence tables."
        )

    print(
        f"Markdown report: {markdown_path}"
    )

    print(
        f"Markdown bytes: "
        f"{markdown_path.stat().st_size:,}"
    )

    print(
        f"DOCX report: {docx_path}"
    )

    print(
        f"DOCX bytes: "
        f"{docx_path.stat().st_size:,}"
    )

    print(
        f"DOCX paragraphs: "
        f"{len(document.paragraphs):,}"
    )

    print(
        f"DOCX tables: "
        f"{len(document.tables):,}"
    )

    print()
    print("=" * 80)
    print(
        "AI report export validation "
        "completed successfully"
    )
    print("=" * 80)


if __name__ == "__main__":
    main()