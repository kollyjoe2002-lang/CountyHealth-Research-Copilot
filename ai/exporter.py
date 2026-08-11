from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ai.formatter import (
    format_evidence_table,
    table_caption,
)
from ai.models import ResearchReport
from ai.report_writer import report_to_markdown


class ReportExportError(RuntimeError):
    """Raised when a report cannot be exported safely."""


def safe_filename_component(
    value: str,
) -> str:
    """
    Convert text into a filesystem-safe filename component.
    """
    cleaned = re.sub(
        r"[^A-Za-z0-9_-]+",
        "_",
        value.strip(),
    )

    cleaned = cleaned.strip("_").lower()

    return cleaned or "countyhealth_report"


def default_report_filename(
    report: ResearchReport,
    extension: str,
) -> str:
    """
    Build a descriptive report filename.
    """
    timestamp = datetime.now().strftime(
        "%Y%m%d_%H%M%S"
    )

    title = safe_filename_component(
        report.title
    )

    extension = extension.lstrip(".")

    return (
        f"{title}_{timestamp}.{extension}"
    )


def export_markdown_bytes(
    report: ResearchReport,
) -> bytes:
    """
    Export a report as UTF-8 Markdown bytes.
    """
    markdown = report_to_markdown(
        report
    )

    return markdown.encode("utf-8")


def save_markdown_report(
    report: ResearchReport,
    output_path: str | Path,
) -> Path:
    """
    Save a report as a Markdown file.
    """
    path = Path(output_path)

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    path.write_bytes(
        export_markdown_bytes(report)
    )

    return path


def _set_document_defaults(
    document: Document,
) -> None:
    """
    Apply consistent document typography and margins.
    """
    section = document.sections[0]

    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10.5)

    for style_name, size in [
        ("Title", 20),
        ("Heading 1", 15),
        ("Heading 2", 12),
    ]:
        style = document.styles[
            style_name
        ]

        style.font.name = "Aptos"
        style.font.size = Pt(size)


def _add_report_header(
    document: Document,
    report: ResearchReport,
) -> None:
    """
    Add the report title and generation metadata.
    """
    title = document.add_heading(
        report.title,
        level=0,
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle = document.add_paragraph(
        "CountyHealth Research Copilot"
    )

    subtitle.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle_run = subtitle.runs[0]
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(11)

    generated = document.add_paragraph(
        "Generated: "
        + datetime.now().strftime(
            "%B %d, %Y at %I:%M %p"
        )
    )

    generated.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    document.add_paragraph()


def _add_bullets(
    document: Document,
    values: list[str],
) -> None:
    """
    Add a list of bullet points.
    """
    for value in values:
        document.add_paragraph(
            value,
            style="List Bullet",
        )


def _clean_cell_value(
    value: Any,
) -> str:
    """
    Convert table-cell values to display text.
    """
    if value is None or pd.isna(value):
        return ""

    if isinstance(value, float):
        return f"{value:,.2f}"

    return str(value)


def _column_width_inches(
    column_name: str,
) -> float:
    """
    Return a practical DOCX width for a formatted evidence column.

    Widths are based on semantic content rather than equal-width
    distribution.
    """
    normalized = (
        column_name
        .strip()
        .casefold()
    )

    if normalized in {
        "county",
        "cause",
        "metric",
    }:
        return 1.90

    if normalized in {
        "rank",
        "county rank",
    }:
        return 0.55

    if normalized == "fips":
        return 0.65

    if normalized == "year":
        return 0.65

    if normalized == "direction":
        return 0.85

    if normalized in {
        "estimate",
        "lower",
        "upper",
    }:
        return 0.75

    if normalized in {
        "yll rate",
        "2000 rate",
        "2019 rate",
    }:
        return 0.80

    if normalized in {
        "national rank",
        "percentile",
        "signed gap",
        "gap magnitude",
        "relative gap",
        "absolute change",
        "percent change",
    }:
        return 0.90

    # Dynamic disparity columns such as:
    # "Non-Latino, Black Rate"
    # "Non-Latino, White Rate"
    if normalized.endswith(" rate"):
        return 1.15

    return 0.85


def _table_requires_landscape(
    dataframe: pd.DataFrame,
) -> bool:
    """
    Determine whether a formatted evidence table should be
    displayed in landscape orientation.

    Wide analytical tables are moved to landscape even when they
    could technically fit on a portrait page, because the wider
    layout substantially improves readability.
    """
    if dataframe.empty:
        return False

    total_width = sum(
        _column_width_inches(
            str(column)
        )
        for column in dataframe.columns
    )

    # In practice, tables approaching six inches become crowded
    # once Word cell padding and wrapping are considered.
    return total_width > 5.8


def _start_landscape_section(
    document: Document,
):
    """
    Start a new landscape section and return it.
    """
    section = document.add_section()

    section.orientation = (
        WD_ORIENT.LANDSCAPE
    )

    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )

    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)

    return section


def _start_portrait_section(
    document: Document,
):
    """
    Start a new portrait section and return it.
    """
    section = document.add_section()

    section.orientation = (
        WD_ORIENT.PORTRAIT
    )

    if (
        section.page_width
        > section.page_height
    ):
        (
            section.page_width,
            section.page_height,
        ) = (
            section.page_height,
            section.page_width,
        )

    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.80)
    section.right_margin = Inches(0.80)

    return section


def _add_dataframe_table(
    document: Document,
    dataframe: pd.DataFrame,
    *,
    maximum_rows: int = 20,
    manage_orientation: bool = True,
) -> bool:
    """
    Add a compact, publication-oriented evidence table
    to the DOCX report.

    Returns True when the formatted table qualifies for
    landscape orientation.
    """
    if dataframe.empty:
        document.add_paragraph(
            "No records were returned."
        )

        return False

    use_landscape = (
        _table_requires_landscape(
            dataframe
        )
    )

    if (
        use_landscape
        and manage_orientation
    ):
        _start_landscape_section(
            document
        )

    display = dataframe.head(
        maximum_rows
    ).copy()

    table = document.add_table(
        rows=1,
        cols=len(display.columns),
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.style = "Table Grid"

    # Keep Word from continuously redistributing
    # widths according to cell contents.
    table.autofit = False

    column_widths = [
        Inches(
            _column_width_inches(
                str(column)
            )
        )
        for column in display.columns
    ]

    # Set the underlying table column widths.
    for index, width in enumerate(
        column_widths
    ):
        table.columns[index].width = (
            width
        )

    header_cells = (
        table.rows[0].cells
    )

    for index, column in enumerate(
        display.columns
    ):
        header_cells[index].width = (
            column_widths[index]
        )

        header_cells[index].text = (
            str(column)
        )

        paragraph = (
            header_cells[index]
            .paragraphs[0]
        )

        paragraph.paragraph_format.space_before = (
            Pt(0)
        )

        paragraph.paragraph_format.space_after = (
            Pt(0)
        )

        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(7.5)

    for row in display.itertuples(
        index=False,
        name=None,
    ):
        cells = table.add_row().cells

        for index, value in enumerate(
            row
        ):
            cells[index].width = (
                column_widths[index]
            )

            cells[index].text = (
                _clean_cell_value(
                    value
                )
            )

            for paragraph in (
                cells[index].paragraphs
            ):
                paragraph.paragraph_format.space_before = (
                    Pt(0)
                )

                paragraph.paragraph_format.space_after = (
                    Pt(0)
                )

                for run in paragraph.runs:
                    run.font.size = Pt(7.5)

    if (
        len(dataframe)
        > maximum_rows
    ):
        note = document.add_paragraph(
            f"Table displays the first "
            f"{maximum_rows:,} of "
            f"{len(dataframe):,} records."
        )

        for run in note.runs:
            run.font.size = Pt(8)
            run.italic = True

    if (
        use_landscape
        and manage_orientation
    ):
        _start_portrait_section(
            document
        )

    return use_landscape


def build_docx_document(
    report: ResearchReport,
    *,
    include_evidence_tables: bool = True,
    evidence_row_limit: int = 20,
) -> Document:
    """
    Build a DOCX document from a structured research report.
    """
    document = Document()

    _set_document_defaults(
        document
    )

    _add_report_header(
        document,
        report,
    )

    document.add_heading(
        "Research Question",
        level=1,
    )

    document.add_paragraph(
        report.question
    )

    document.add_heading(
        "Executive Summary",
        level=1,
    )

    document.add_paragraph(
        report.executive_summary
    )

    document.add_heading(
        "Methods",
        level=1,
    )

    document.add_paragraph(
        report.methods
    )

    document.add_heading(
        "Key Findings",
        level=1,
    )

    _add_bullets(
        document,
        report.findings,
    )

    document.add_heading(
        "Limitations",
        level=1,
    )

    _add_bullets(
        document,
        report.limitations,
    )

    if report.evidence.warnings:
        document.add_heading(
            "Data Warnings",
            level=1,
        )

        _add_bullets(
            document,
            report.evidence.warnings,
        )

    document.add_heading(
        "Analytical Provenance",
        level=1,
    )

    document.add_paragraph(
        "This report was generated from validated "
        "CountyHealth Research Copilot analytical "
        "functions. Each evidence item records the "
        "source function and execution parameters."
    )

    for index, item in enumerate(
        report.evidence.items,
        start=1,
    ):
        document.add_heading(
            f"Evidence {index}: {item.title}",
            level=2,
        )

        document.add_paragraph(
            f"Source function: "
            f"{item.source_function}"
        )

        parameters_text = ", ".join(
            f"{key}={value}"
            for key, value in (
                item.parameters.items()
            )
        )

        document.add_paragraph(
            "Parameters: "
            + (
                parameters_text
                if parameters_text
                else "None"
            )
        )

        if item.interpretation_note:
            document.add_paragraph(
                item.interpretation_note
            )

        if (
            include_evidence_tables
            and isinstance(
                item.data,
                pd.DataFrame,
            )
        ):
            formatted_table = (
                format_evidence_table(
                    item.source_function,
                    item.data,
                    context=(
                        report.evidence.context
                    ),
                )
            )

            caption_text = (
                table_caption(
                    item.source_function,
                    item.title,
                )
            )

            use_landscape = (
                _table_requires_landscape(
                    formatted_table
                )
            )

            # Start the landscape section BEFORE the caption
            # so the caption and table remain together.
            if use_landscape:
                _start_landscape_section(
                    document
                )

            caption = (
                document.add_paragraph()
            )

            caption_run = caption.add_run(
                f"Table {index}. "
                f"{caption_text}"
            )

            caption_run.bold = True
            caption_run.font.size = Pt(9)

            _add_dataframe_table(
                document,
                formatted_table,
                maximum_rows=(
                    evidence_row_limit
                ),
                manage_orientation=False,
            )

            # Return subsequent report content to portrait.
            if use_landscape:
                _start_portrait_section(
                    document
                )

    document.add_heading(
        "Data Source and Interpretation Note",
        level=1,
    )

    document.add_paragraph(
        "The report is based on IHME-derived "
        "county-level estimates incorporated into "
        "the local CountyHealth DuckDB warehouse. "
        "Users should consult the original IHME "
        "documentation for authoritative methodological "
        "and citation guidance."
    )

    return document


def export_docx_bytes(
    report: ResearchReport,
    *,
    include_evidence_tables: bool = True,
    evidence_row_limit: int = 20,
) -> bytes:
    """
    Export a report as DOCX bytes.
    """
    document = build_docx_document(
        report,
        include_evidence_tables=(
            include_evidence_tables
        ),
        evidence_row_limit=(
            evidence_row_limit
        ),
    )

    buffer = BytesIO()

    document.save(
        buffer
    )

    buffer.seek(0)

    return buffer.getvalue()


def save_docx_report(
    report: ResearchReport,
    output_path: str | Path,
    *,
    include_evidence_tables: bool = True,
    evidence_row_limit: int = 20,
) -> Path:
    """
    Save a structured research report as a DOCX file.
    """
    path = Path(
        output_path
    )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    path.write_bytes(
        export_docx_bytes(
            report,
            include_evidence_tables=(
                include_evidence_tables
            ),
            evidence_row_limit=(
                evidence_row_limit
            ),
        )
    )

    return path